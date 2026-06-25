#!/usr/bin/env python3
"""Inspect core structure and style metadata from a .docx file."""

from __future__ import annotations

import argparse
import json
from collections import Counter
from pathlib import Path

from docx import Document


def inspect_docx(path: Path) -> dict:
    doc = Document(path)

    paragraph_styles = Counter(p.style.name if p.style else "(none)" for p in doc.paragraphs)
    table_styles = Counter(t.style.name if t.style else "(none)" for t in doc.tables)

    headings = []
    nonempty_paragraphs = 0
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            nonempty_paragraphs += 1
        style = paragraph.style.name if paragraph.style else ""
        if style.lower().startswith("heading") or style.startswith("标题"):
            headings.append({"index": index, "style": style, "text": text[:120]})

    sections = []
    for index, section in enumerate(doc.sections, start=1):
        sections.append(
            {
                "index": index,
                "page_width_twips": section.page_width.twips,
                "page_height_twips": section.page_height.twips,
                "top_margin_twips": section.top_margin.twips,
                "bottom_margin_twips": section.bottom_margin.twips,
                "left_margin_twips": section.left_margin.twips,
                "right_margin_twips": section.right_margin.twips,
                "header_distance_twips": section.header_distance.twips,
                "footer_distance_twips": section.footer_distance.twips,
            }
        )

    return {
        "path": str(path),
        "paragraph_count": len(doc.paragraphs),
        "nonempty_paragraph_count": nonempty_paragraphs,
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "paragraph_styles": dict(paragraph_styles.most_common()),
        "table_styles": dict(table_styles.most_common()),
        "headings": headings,
        "sections": sections,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Inspect a .docx file.")
    parser.add_argument("docx_path", type=Path)
    parser.add_argument("--json", action="store_true", help="Emit compact JSON only.")
    args = parser.parse_args()

    info = inspect_docx(args.docx_path)
    if args.json:
        print(json.dumps(info, ensure_ascii=False, indent=2))
        return

    print(f"File: {info['path']}")
    print(f"Paragraphs: {info['paragraph_count']} ({info['nonempty_paragraph_count']} non-empty)")
    print(f"Tables: {info['table_count']}")
    print(f"Sections: {info['section_count']}")
    print("\nParagraph styles:")
    for name, count in info["paragraph_styles"].items():
        print(f"  {name}: {count}")
    if info["table_styles"]:
        print("\nTable styles:")
        for name, count in info["table_styles"].items():
            print(f"  {name}: {count}")
    if info["headings"]:
        print("\nHeadings:")
        for heading in info["headings"]:
            print(f"  [{heading['index']}] {heading['style']}: {heading['text']}")
    print("\nSections:")
    for section in info["sections"]:
        print(f"  Section {section['index']}: {section}")


if __name__ == "__main__":
    main()
